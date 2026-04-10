# -*- coding: utf-8 -*-
"""
Руководство к файлу llm_service.py
===================================

Назначение:
    Единый асинхронный LLM-сервис (глобальный модуль), покрывающий все
    use-cases из USE_CASES.md / LLM_SERVICE.md (UC-1 … UC-20):
      UC-1  — RAG (hybrid: dense + sparse) + интернет + tool calling + стриминг + JSON repair
      UC-2  — Инжест документов (PDF/DOCX/DOC/TXT/MD/HTML) без векторизации
      UC-3  — Cloud.ru ASR → LLM-ответ
      UC-4  — Streaming-чат с промежуточными токенами
      UC-5  — Tool calling с лимитами и цепочками
      UC-20 — Безопасная транскрибация больших аудио (сжатие + нарезка + очередь)

Единый контракт:
    UserInput  — текст / audio_ref / files + флаги (mode, USE_INTERNET,
                 ENABLE_TOOLS, ENABLE_STREAMING, ENABLE_ASR).
    RequestContext — request_id, user_id, RAG/tool параметры, rag_filters.

История переписки (conversation history):
    AI-сервис НЕ хранит историю — она живёт в БД проекта-потребителя.
    Flow:
      1. API получает запрос с conversation_id (или без).
      2. Оркестратор проекта загружает историю из БД → List[ChatMessage].
      3. Передаёт историю через UserInput.history в AI-сервис.
      4. ContextBuilder вставляет историю между system-промптом и RAG/doc контекстом.
      5. После ответа LLM оркестратор сохраняет user + assistant сообщения в БД.
    Конвертация ChatMessage → LLMMessage (OpenAI-формат): ContextBuilder._chat_message_to_llm().

ASR автоподготовка (UC-3 + UC-20):
    Если файл не проходит безопасный порог размера или формат неподходящий,
    ASRClient автоматически делает fallback-подготовку через ffmpeg:
    1) сжатие в Opus/AAC (без PCM-раздувания),
    2) при необходимости — нарезка на чанки < лимита,
    3) последовательная отправка чанков с сохранением порядка.
    Это позволяет избегать 413 Payload Too Large и детально диагностировать
    проблемы по шагам логирования.

Pipeline оркестратора (ChatOrchestrator):
    ШАГ 0  — Инициализация + ASR (UC-3).
    ШАГ 1  — (опционально) подготовка данных RAG/ингеста.
    ШАГ 2  — Инжест документов (UC-2).
    ШАГ 3  — RAG + Internet параллельно (UC-1):
             3.1 RagQueryRewriter → 3.3a dense embed + 3.3b sparse embed
             → 3.3c hybrid search (RRF fusion) с фильтрацией
             3.2 CrawlerQueryRewriter → 3.4 crawler4ai.search
    ШАГ 4  — ContextBuilder.build_input (сниппеты + документы → LLMRequest).
    ШАГ 5  — Tool calls цикл до tool_call_limit (UC-5).
    ШАГ 6  — LLM ответ + подсчёт токенов (prompt/completion/total) + JSON repair.
    ШАГ 7  — StreamEvent(step="done").

Логирование:
    Все шаги логируются в формате «ШАГ N. <описание> — ОТПРАВЛЯЕМ / ОЖИДАЕМ / УСПЕХ / ОШИБКА».

Qdrant DTO (конфигурация коллекций, sparse-вектора, гибридный поиск):
    SparseVectorData         — Pydantic-обёртка sparse-вектора (indices + values).
                              Валидация (длины, уникальность indices), свойства nnz/is_empty,
                              метод to_qdrant() → qdrant_client.models.SparseVector.
    QdrantCollectionConfig   — Конфигурация коллекции: named dense vector (size, distance,
                              on_disk), named sparse vector (IDF modifier, on_disk),
                              payload index fields, HNSW/optimizer параметры.
    HybridSearchConfig       — Конфигурация hybrid search: тип fusion (rrf/dbsf),
                              множители prefetch limit, score_threshold, with_payload/vectors.

Векторизация и индексация (Qdrant — dense + sparse неразрывно):
    CloudRuEmbeddingClient  — dense embedding через OpenAI-совместимый API Cloud.ru
                              (POST /v1/embeddings, батч-обработка).
    SparseEmbeddingClient   — sparse embedding через sentence-transformers SparseEncoder
                              (модель opensearch-neural-sparse-encoding-multilingual-v1).
                              Возвращает SparseVectorData. Единый парсер _parse_raw_embedding
                              обрабатывает все форматы: dict, scipy, torch, numpy, SparseEncoding.
    QdrantVectorStore        — Qdrant коллекция с named vectors:
                              "dense" (VectorParams, COSINE) + "sparse" (SparseVectorParams, IDF).
                              Создание: collection_exists() (v1.8+), payload indexes сразу.
                              Hybrid search: prefetch dense + prefetch sparse → RRF/DBSF fusion.
                              Утилиты: delete_collection, get_collection_info, count_points.
                              Фильтрация по payload: user_id, document_name, tags, custom_data.*.
    ingest_and_index()       — bulk-индексация: файлы → ingest → chunk →
                              dense embed + sparse embed → upsert (неразрывно)
                              в Qdrant с метаданными из IndexMetadata.

Метаданные и фильтрация:
    IndexMetadata            — DTO для передачи метаданных при индексации:
                              document_name, user_id, tags, custom_data (dict).
                              Все поля сохраняются в payload точки Qdrant.
    ctx.rag_filters          — dict фильтров для поиска:
                              {"user_id": "...", "document_name": "...",
                               "tags": [...], "custom_data.<key>": value}.
                              Автоматически конвертируются в Qdrant Filter (must/AND).

Зависимости ядра (без FastAPI):
    pydantic, openai (AsyncOpenAI), httpx, asyncio, json, logging.
    Опциональные (lazy import): sentence_transformers (SparseEncoder),
    qdrant_client, fitz (PyMuPDF), python-docx, html.parser.

Модульность (фасад + вынесенные подсистемы):
    Для сохранения обратной совместимости импортов use-cases файл остаётся
    фасадом AI.llm_service, а тяжёлые подсистемы вынесены в отдельные модули:
      - AI.llm_asr                (TranscriptSegment, Transcript, ASRClient)
      - AI.llm_qdrant             (Sparse/Qdrant DTO, Embedding, VectorStore,
                                   ingest_and_index, Retriever)
      - AI.llm_webcrawler         (CrawlerQueryRewriter, CrawlerClient)
      - AI.llm_semantic_scholar   (S2Client, S2SearchFilter, S2FieldInference,
                                   S2_* константы)

    ВНИМАНИЕ:
      - Публичный API (имена в __all__) сохранён.
      - Use-cases продолжают импортировать только из AI.llm_service.
      - Внизу файла выполнен явный реэкспорт вынесенных сущностей.

"""

# Руководство к файлу:
# llm_service.py — фасад AI-интеграций и оркестрации MyPerplexity.
# Содержит модели запросов/ответов, LLM/RAG/WebCrawler/ASR интеграции, ChatOrchestrator и фабрики создания клиентов.
# При изменениях в pipeline важно сохранять совместимость публичных импортов, подробное шаговое логирование и source-of-truth для orchestration flow.
# Internet/web pipeline обязан запускаться независимо от RAG-режима, если включён флаг use_internet и доступен текст запроса.

from __future__ import annotations

import asyncio
import hashlib
import json
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
import uuid
from pathlib import Path
from datetime import datetime
from typing import (
    Any,
    AsyncIterator,
    Awaitable,
    Callable,
    Dict,
    List,
    Optional,
    Sequence,
    Tuple,
)

from dataclasses import dataclass
from pydantic import BaseModel, Field, ValidationError
from typing_extensions import Literal

import httpx
from openai import AsyncOpenAI

logger = logging.getLogger(__name__)


# ============================================================================
# 1. Базовые DTO / Pydantic-схемы
# ============================================================================


class TokenUsage(BaseModel):
    """Подсчёт токенов (prompt / completion / total)."""

    prompt_tokens: Optional[int] = None
    completion_tokens: Optional[int] = None
    total_tokens: Optional[int] = None


class RequestContext(BaseModel):
    """Контекст запроса: идентификаторы + флаги поведения + параметры RAG/tools."""

    request_id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    conversation_id: Optional[str] = None
    user_id: Optional[str] = None

    # --- Флаги поведения (синхронизируются из UserInput) ---
    mode: Literal["chat", "rag_tool", "rag_qa"] = "chat"
    use_internet: bool = False
    enable_tools: bool = False
    enable_streaming: bool = True
    enable_asr: bool = False

    # --- Параметры RAG ---
    rag_top_k: int = 6
    rag_rewrite_count: int = 3
    rag_filters: Optional[Dict[str, Any]] = None

    # --- Параметры tools ---
    tool_whitelist: Optional[List[str]] = None
    tool_call_limit: int = 3
    tool_timeout: float = 5.0

    # --- Общие ---
    timeout: Optional[float] = None
    trace_id: Optional[str] = None
    metadata: Dict[str, Any] = Field(default_factory=dict)


class ToolCall(BaseModel):
    id: str
    name: str
    arguments: Dict[str, Any]


class ToolResult(BaseModel):
    id: str
    name: str
    output: Dict[str, Any]
    is_error: bool = False
    error_message: Optional[str] = None


class ChatMessage(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    role: Literal["system", "user", "assistant", "tool"]
    content: str
    created_at: datetime = Field(default_factory=datetime.utcnow)
    tokens: Optional[int] = None
    tool_calls: Optional[List[ToolCall]] = None
    tool_call_id: Optional[str] = None
    metadata: Dict[str, Any] = Field(default_factory=dict)


class FileRef(BaseModel):
    """Ссылка на файл (для инжеста документов, UC-2)."""

    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    path: Optional[str] = None
    url: Optional[str] = None
    mime_type: Optional[str] = None
    original_name: Optional[str] = None
    size_bytes: Optional[int] = None
    metadata: Dict[str, Any] = Field(default_factory=dict)


class UserInput(BaseModel):
    """Единый контракт пользовательского ввода для всех UC.

    Флаги: mode, use_internet, enable_tools, enable_streaming, enable_asr.
    """

    conversation_id: Optional[str] = None
    user_id: Optional[str] = None

    text: Optional[str] = None
    audio_ref: Optional[str] = None
    files: Optional[List[FileRef]] = None

    mode: Literal["chat", "rag_tool", "rag_qa"] = "chat"
    use_internet: bool = False
    enable_tools: bool = False
    enable_streaming: bool = True
    enable_asr: bool = False

    language: Optional[str] = None
    metadata: Dict[str, Any] = Field(default_factory=dict)

    # --- История переписки (загружается оркестратором из БД) ---
    history: Optional[List["ChatMessage"]] = None


class LLMMessage(BaseModel):
    role: Literal["system", "user", "assistant", "tool"]
    content: str
    name: Optional[str] = None
    tool_call_id: Optional[str] = None
    tool_calls: Optional[List[Dict[str, Any]]] = None


class LLMRequest(BaseModel):
    messages: List[LLMMessage]
    model: str
    temperature: Optional[float] = None
    max_output_tokens: Optional[int] = None
    tools: Optional[List[Dict[str, Any]]] = None
    tool_choice: Optional[Any] = None
    response_format: Optional[Dict[str, Any]] = None
    metadata: Dict[str, Any] = Field(default_factory=dict)


class LLMResponse(BaseModel):
    content: str
    tool_calls: List[ToolCall] = Field(default_factory=list)
    usage: Optional[TokenUsage] = None
    raw_response: Dict[str, Any] = Field(default_factory=dict)


class StreamEvent(BaseModel):
    """Событие стрима (UC-4).

    kind: init | token_delta | tool_call | tool_result | rag_progress |
          internet_progress | final | done | error.
    """

    kind: Literal[
        "init", "token_delta", "tool_call", "tool_result",
        "rag_progress", "internet_progress", "final", "done", "error",
    ]
    step: Optional[str] = None
    data: Dict[str, Any] = Field(default_factory=dict)


class DocumentChunk(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    text: str
    source_id: str
    page: Optional[int] = None
    offset: Optional[int] = None
    checksum: Optional[str] = None
    metadata: Dict[str, Any] = Field(default_factory=dict)


class IndexMetadata(BaseModel):
    """Метаданные для индексации чанков в Qdrant.

    Сохраняются в payload каждой точки и доступны для фильтрации при поиске.
    """

    document_name: Optional[str] = None
    user_id: Optional[str] = None
    tags: List[str] = Field(default_factory=list)
    custom_data: Dict[str, Any] = Field(default_factory=dict)


class Snippet(BaseModel):
    text: str
    source_id: str
    score: float
    metadata: Dict[str, Any] = Field(default_factory=dict)


# ============================================================================
# 1b. Qdrant/ASR DTO вынесены в модули (см. реэкспорт в секции 16c)
# ============================================================================


class OrchestratorResult(BaseModel):
    """Финальный результат ChatOrchestrator.handle_user_input."""

    response_text: str
    sources: List[Snippet] = Field(default_factory=list)
    tool_calls: List[ToolCall] = Field(default_factory=list)
    tokens_used: Optional[TokenUsage] = None
    context_used: Dict[str, Any] = Field(default_factory=dict)


# ============================================================================
# 3. ContextBuilder — сборка LLMRequest из RAG + интернет + документов
# ============================================================================


class ContextBuilder:
    """Собирает LLMRequest из RAG + интернет + документов."""

    def __init__(self, default_model: str, system_prompt: Optional[str] = None) -> None:
        self._default_model = default_model
        self._system_prompt = system_prompt or (
            "Вы — полезный ассистент. Отвечайте точно и по делу, "
            "используя предоставленный контекст."
        )

    @staticmethod
    def _chat_message_to_llm(msg: "ChatMessage") -> LLMMessage:
        """Конвертация ChatMessage (из БД оркестратора) → LLMMessage (OpenAI-формат).

        Маппинг:
          - role, content → напрямую
          - tool_calls → list[dict] в формате OpenAI function calling
          - tool_call_id → для role="tool"
        """
        tool_calls_raw: Optional[List[Dict[str, Any]]] = None
        if msg.tool_calls:
            tool_calls_raw = [
                {
                    "id": tc.id,
                    "type": "function",
                    "function": {
                        "name": tc.name,
                        "arguments": json.dumps(tc.arguments, ensure_ascii=False),
                    },
                }
                for tc in msg.tool_calls
            ]
        return LLMMessage(
            role=msg.role,
            content=msg.content,
            tool_call_id=msg.tool_call_id,
            tool_calls=tool_calls_raw,
        )

    async def build_input(
        self,
        user_input: UserInput,
        rag_snippets: Optional[List[Snippet]],
        internet_snippets: Optional[List[Snippet]],
        document_chunks: Optional[List[DocumentChunk]],
        ctx: RequestContext,
        history: Optional[List["ChatMessage"]] = None,
    ) -> LLMRequest:
        logger.info(
            "ШАГ 4. ContextBuilder.build_input — собираем промпт: "
            "request_id=%s, conversation_id=%s, history_len=%d",
            ctx.request_id, ctx.conversation_id,
            len(history) if history else 0,
        )

        messages: List[LLMMessage] = []
        messages.append(LLMMessage(role="system", content=self._system_prompt))

        # --- История переписки (между system-промптом и RAG/doc контекстом) ---
        if history:
            logger.info(
                "ШАГ 4.1. Вставляем историю переписки: %d сообщений",
                len(history),
            )
            for msg in history:
                messages.append(self._chat_message_to_llm(msg))

        if rag_snippets:
            parts = [
                f"[RAG #{i + 1} score={s.score:.3f}] {s.text}"
                for i, s in enumerate(rag_snippets)
            ]
            messages.append(LLMMessage(
                role="system",
                content=(
                    "Релевантные фрагменты из базы знаний "
                    "(не ссылайтесь на нумерацию напрямую):\n\n"
                    + "\n\n".join(parts)
                ),
            ))

        if internet_snippets:
            parts = [f"[WEB #{i + 1}] {s.text}" for i, s in enumerate(internet_snippets)]
            messages.append(LLMMessage(
                role="system",
                content="Результаты из интернета:\n\n" + "\n\n".join(parts),
            ))

        if document_chunks:
            parts = [
                f"[DOC {c.source_id} p.{c.page or '?'}] {c.text}"
                for c in document_chunks
            ]
            messages.append(LLMMessage(
                role="system",
                content="Содержимое загруженных документов:\n\n" + "\n\n".join(parts),
            ))

        if user_input.text:
            messages.append(LLMMessage(role="user", content=user_input.text))

        logger.info(
            "ШАГ 4. ContextBuilder.build_input — УСПЕХ: messages=%d",
            len(messages),
        )
        return LLMRequest(messages=messages, model=self._default_model)


# ============================================================================
# 4. JsonRepairLLM + StrictOutputParser
# ============================================================================


class JsonRepairLLM:
    """Пытается починить невалидный JSON из ответа LLM."""

    async def repair(self, raw_text: str, ctx: RequestContext) -> str:
        logger.info(
            "ШАГ JSON-REPAIR. Попытка починить JSON: request_id=%s, len=%d",
            ctx.request_id, len(raw_text),
        )
        text = raw_text.strip()

        if text.startswith("```"):
            lines = text.split("\n")
            lines = [ln for ln in lines if not ln.strip().startswith("```")]
            text = "\n".join(lines)

        json_match = re.search(r"\{[\s\S]*\}", text)
        if json_match:
            text = json_match.group(0)

        text = re.sub(r",\s*([\]}])", r"\1", text)

        try:
            json.loads(text)
            logger.info("ШАГ JSON-REPAIR. УСПЕХ")
            return text
        except json.JSONDecodeError:
            logger.warning(
                "ШАГ JSON-REPAIR. Не удалось починить JSON, возвращаем оригинал"
            )
            return raw_text


class StrictOutputParser:
    """Строгая валидация вывода LLM через Pydantic-схемы + авто-ремонт JSON."""

    def __init__(self, json_repair: Optional[JsonRepairLLM] = None) -> None:
        self._repair = json_repair or JsonRepairLLM()

    async def parse_json(
        self, text: str, model: type[BaseModel], ctx: RequestContext,
    ) -> BaseModel:
        logger.info(
            "ШАГ PARSE. Парсим JSON через Pydantic: request_id=%s, schema=%s",
            ctx.request_id, model.__name__,
        )
        try:
            data = json.loads(text)
        except json.JSONDecodeError:
            repaired = await self._repair.repair(text, ctx)
            try:
                data = json.loads(repaired)
            except json.JSONDecodeError as exc:
                logger.error(
                    "ШАГ PARSE. ОШИБКА — JSON невалиден даже после ремонта: %s", exc,
                )
                raise

        try:
            if hasattr(model, "model_validate"):
                return model.model_validate(data)
            return model.parse_obj(data)  # type: ignore[union-attr]
        except ValidationError as exc:
            logger.error("ШАГ PARSE. ОШИБКА — Pydantic validation: %s", exc)
            raise


# ============================================================================
# 5. Tool system (UC-5)
# ============================================================================


class ToolSpec(BaseModel):
    """Описание одного tool: имя, описание, Pydantic-схемы, таймаут."""

    name: str
    description: str
    args_schema: type[BaseModel]
    result_schema: Optional[type[BaseModel]] = None
    timeout: float = 5.0
    metadata: Dict[str, Any] = Field(default_factory=dict)

    class Config:
        arbitrary_types_allowed = True


class ToolRequestContext(BaseModel):
    request_id: str
    conversation_id: Optional[str] = None
    user_id: Optional[str] = None
    tool_name: str
    trace_id: Optional[str] = None
    metadata: Dict[str, Any] = Field(default_factory=dict)


ToolHandler = Callable[[BaseModel, ToolRequestContext], Awaitable[Optional[BaseModel]]]


class ToolRegistry:
    """Реестр tools. Хранит ToolSpec + хендлеры."""

    def __init__(self) -> None:
        self._tools: Dict[str, ToolSpec] = {}
        self._handlers: Dict[str, ToolHandler] = {}

    async def register(self, tool: ToolSpec, handler: ToolHandler) -> None:
        logger.info("ToolRegistry.register: tool=%s", tool.name)
        self._tools[tool.name] = tool
        self._handlers[tool.name] = handler

    def get_openai_tools(
        self, whitelist: Optional[List[str]] = None,
    ) -> List[Dict[str, Any]]:
        """Сериализация ToolSpec → формат tools для OpenAI API."""
        result: List[Dict[str, Any]] = []
        for name, spec in self._tools.items():
            if whitelist and name not in whitelist:
                continue
            schema = spec.args_schema.model_json_schema()
            result.append({
                "type": "function",
                "function": {
                    "name": spec.name,
                    "description": spec.description,
                    "parameters": schema,
                },
            })
        return result

    def get_spec(self, name: str) -> ToolSpec:
        if name not in self._tools:
            raise KeyError(f"Неизвестный tool: {name}")
        return self._tools[name]

    def get_handler(self, name: str) -> ToolHandler:
        if name not in self._handlers:
            raise KeyError(f"Нет хендлера для tool: {name}")
        return self._handlers[name]


class ToolExecutor:
    """Исполняет tool calls с валидацией аргументов, таймаутами и лимитами."""

    def __init__(self, registry: ToolRegistry, concurrency_limit: int = 8) -> None:
        self._registry = registry
        self._semaphore = asyncio.Semaphore(concurrency_limit)

    async def execute(self, call: ToolCall, ctx: RequestContext) -> ToolResult:
        logger.info(
            "ШАГ TOOL. Вызов tool=%s id=%s: request_id=%s — ОТПРАВЛЯЕМ",
            call.name, call.id, ctx.request_id,
        )
        async with self._semaphore:
            spec = self._registry.get_spec(call.name)
            handler = self._registry.get_handler(call.name)

            tool_ctx = ToolRequestContext(
                request_id=ctx.request_id,
                conversation_id=ctx.conversation_id,
                user_id=ctx.user_id,
                tool_name=call.name,
                trace_id=ctx.trace_id,
            )

            try:
                args_model = spec.args_schema(**call.arguments)
                result_model = await asyncio.wait_for(
                    handler(args_model, tool_ctx),
                    timeout=spec.timeout,
                )
                output = result_model.model_dump() if result_model else {}
                logger.info("ШАГ TOOL. tool=%s — УСПЕХ", call.name)
                return ToolResult(id=call.id, name=call.name, output=output)

            except asyncio.TimeoutError:
                msg = f"Timeout after {spec.timeout}s"
                logger.error("ШАГ TOOL. tool=%s — ОШИБКА: %s", call.name, msg)
                return ToolResult(
                    id=call.id, name=call.name, output={},
                    is_error=True, error_message=msg,
                )
            except ValidationError as exc:
                logger.error(
                    "ШАГ TOOL. tool=%s — ОШИБКА валидации аргументов: %s",
                    call.name, exc,
                )
                return ToolResult(
                    id=call.id, name=call.name, output={},
                    is_error=True, error_message=f"Validation error: {exc}",
                )
            except Exception as exc:
                logger.error("ШАГ TOOL. tool=%s — ОШИБКА: %s", call.name, exc)
                return ToolResult(
                    id=call.id, name=call.name, output={},
                    is_error=True, error_message=str(exc),
                )

    async def execute_many(
        self, calls: Sequence[ToolCall], ctx: RequestContext,
    ) -> List[ToolResult]:
        tasks = [self.execute(c, ctx) for c in calls]
        return list(await asyncio.gather(*tasks))


# ============================================================================
# 6. Инжест документов (UC-2): Chunker + DocumentIngestor
# ============================================================================


class Chunker:
    """Разбивает текст на чанки по токенам с фиксированным перекрытием."""

    def __init__(
        self,
        chunk_size_tokens: int = 512,
        chunk_overlap_tokens: int = 128,
        tokenizer_name: str = "cl100k_base",
    ) -> None:
        self._chunk_size = chunk_size_tokens
        self._overlap = chunk_overlap_tokens
        self._tokenizer_name = tokenizer_name
        self._encoding = None
        self._tokenizer_kind: str = "fallback"
        try:
            import tiktoken  # type: ignore

            self._encoding = tiktoken.get_encoding(self._tokenizer_name)
            self._tokenizer_kind = "tiktoken"
        except Exception:
            # Fallback: простое whitespace-деление
            self._encoding = None
            self._tokenizer_kind = "whitespace"

    def _tokenize(self, text: str) -> List[int | str]:
        if self._tokenizer_kind == "tiktoken" and self._encoding is not None:
            return self._encoding.encode(text)
        return text.split()

    def _decode(self, tokens: List[int | str]) -> str:
        if self._tokenizer_kind == "tiktoken" and self._encoding is not None:
            return self._encoding.decode(tokens)  # type: ignore[arg-type]
        return " ".join(tokens)  # type: ignore[list-item]

    async def split(
        self, text: str, source_id: str, page: Optional[int] = None,
    ) -> List[DocumentChunk]:
        logger.info(
            "ШАГ CHUNK. Разбиваем текст на чанки по токенам: source_id=%s, len=%d, tokenizer=%s",
            source_id, len(text), self._tokenizer_kind,
        )
        if not text.strip():
            return []

        tokens = self._tokenize(text)
        if not tokens:
            return []

        chunks: List[DocumentChunk] = []
        start = 0
        while start < len(tokens):
            end = start + self._chunk_size
            token_slice = tokens[start:end]
            chunk_text = self._decode(token_slice)
            if chunk_text.strip():
                # offset считаем в "токеновых шагах" (start) — пригодно для поиска
                checksum = hashlib.md5(chunk_text.encode()).hexdigest()
                chunks.append(DocumentChunk(
                    text=chunk_text,
                    source_id=source_id,
                    page=page,
                    offset=start,
                    checksum=checksum,
                    metadata={"offset_tokens": start},
                ))
            start = end - self._overlap if end < len(tokens) else end

        logger.info("ШАГ CHUNK. УСПЕХ: %d чанков", len(chunks))
        return chunks


class DocumentIngestor:
    """Парсит PDF / DOCX / DOC / TXT / MD / HTML → список DocumentChunk."""

    # Поддерживаемые расширения для инжеста
    SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".doc", ".html", ".htm", ".txt", ".md")

    def __init__(self, chunker: Optional[Chunker] = None) -> None:
        self._chunker = chunker or Chunker()

    async def ingest(self, file: FileRef, ctx: RequestContext) -> List[DocumentChunk]:
        logger.info(
            "ШАГ INGEST. Инжест документа: request_id=%s, file_id=%s, "
            "mime=%s, path=%s",
            ctx.request_id, file.id, file.mime_type, file.path,
        )
        if not file.path:
            raise ValueError("DocumentIngestor ожидает локальный path")

        ext = file.path.lower().rsplit(".", 1)[-1] if "." in file.path else ""
        mime = (file.mime_type or "").lower()

        try:
            if "pdf" in mime or ext == "pdf":
                pages = await self._parse_pdf(file.path)
            elif ext == "docx" or "officedocument.wordprocessingml" in mime:
                pages = await self._parse_docx(file.path)
            elif ext == "doc" or "msword" in mime:
                pages = await self._parse_docx(file.path)
            elif ext in ("html", "htm") or "html" in mime:
                pages = await self._parse_html(file.path)
            else:
                pages = await self._parse_text(file.path)

            all_chunks: List[DocumentChunk] = []
            for page_num, page_text in pages:
                chunks = await self._chunker.split(
                    page_text, source_id=file.id, page=page_num,
                )
                all_chunks.extend(chunks)

            if not all_chunks:
                logger.warning(
                    "ШАГ INGEST. Парсер вернул пустой текст: file_id=%s", file.id,
                )
            else:
                logger.info(
                    "ШАГ INGEST. УСПЕХ: file_id=%s, chunks=%d",
                    file.id, len(all_chunks),
                )
            return all_chunks

        except Exception as exc:
            logger.error(
                "ШАГ INGEST. ОШИБКА: file_id=%s, error=%s", file.id, exc,
            )
            raise

    async def _parse_pdf(self, path: str) -> List[Tuple[Optional[int], str]]:
        import fitz  # PyMuPDF — lazy import

        result: List[Tuple[Optional[int], str]] = []
        doc = fitz.open(path)
        try:
            for i in range(len(doc)):
                page = doc.load_page(i)
                text = page.get_text("text")
                if text.strip():
                    result.append((i + 1, text))
        finally:
            doc.close()
        return result

    async def _parse_docx(self, path: str) -> List[Tuple[Optional[int], str]]:
        """Парсинг DOCX с разбивкой по page breaks (w:br type='page').

        Если в документе есть явные разрывы страниц — текст разбивается
        на страницы. Если разрывов нет — возвращается единый блок (page=1).
        """
        import docx as python_docx  # lazy import
        from lxml import etree

        document = python_docx.Document(path)
        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        pages: List[Tuple[Optional[int], str]] = []
        current_parts: List[str] = []
        page_num = 1

        for paragraph in document.paragraphs:
            # Проверяем наличие page break в runs параграфа
            has_page_break = False
            for run in paragraph.runs:
                br_elements = run._element.findall(".//w:br", nsmap)
                for br in br_elements:
                    if br.get(f"{{{nsmap['w']}}}type") == "page":
                        has_page_break = True
                        break
                if has_page_break:
                    break

            # Также проверяем pageBreakBefore в стиле параграфа
            pPr = paragraph._element.find("w:pPr", nsmap)
            if pPr is not None:
                pb_before = pPr.find("w:pageBreakBefore", nsmap)
                if pb_before is not None:
                    has_page_break = True

            if has_page_break and current_parts:
                page_text = "\n".join(current_parts)
                if page_text.strip():
                    pages.append((page_num, page_text))
                current_parts = []
                page_num += 1

            if paragraph.text.strip():
                current_parts.append(paragraph.text)

        # Последняя страница
        if current_parts:
            page_text = "\n".join(current_parts)
            if page_text.strip():
                pages.append((page_num, page_text))

        if not pages:
            # Fallback: весь текст как одна страница
            text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
            return [(1, text)] if text.strip() else []

        logger.info(
            "ШАГ INGEST (DOCX). Разбито на %d страниц(ы) по page breaks",
            len(pages),
        )
        return pages

    async def _parse_html(self, path: str) -> List[Tuple[Optional[int], str]]:
        from html.parser import HTMLParser

        class _Stripper(HTMLParser):
            def __init__(self) -> None:
                super().__init__()
                self.parts: List[str] = []

            def handle_data(self, data: str) -> None:
                self.parts.append(data)

        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        stripper = _Stripper()
        stripper.feed(content)
        text = " ".join(stripper.parts)
        return [(None, text)] if text.strip() else []

    async def _parse_text(self, path: str) -> List[Tuple[Optional[int], str]]:
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()
        return [(None, content)] if content.strip() else []

    @staticmethod
    def scan_directory(
        directory_path: str,
        extensions: Optional[Tuple[str, ...]] = None,
        recursive: bool = False,
    ) -> List[str]:
        """ШАГ 1. Получение списка файлов для индексации.

        Функция сканирует директорию и возвращает список путей к файлам
        с поддерживаемыми расширениями.

        Args:
            directory_path: Путь к директории с документами
            extensions: Кортеж расширений (по умолчанию SUPPORTED_EXTENSIONS)
            recursive: Рекурсивный обход поддиректорий

        Returns:
            Список абсолютных путей к файлам (отсортирован)

        Логирование:
            «ШАГ 1. Получили N файлов для индексации — УСПЕХ»
        """
        logger = logging.getLogger(__name__)
        logger.info("ШАГ 1. Получение списка файлов из %s", directory_path)

        files: List[str] = []
        dir_path = Path(directory_path)
        ext_tuple = extensions or DocumentIngestor.SUPPORTED_EXTENSIONS

        if not dir_path.exists():
            logger.error("ШАГ 1. ОШИБКА: Директория не найдена: %s", directory_path)
            return files

        if not dir_path.is_dir():
            logger.error("ШАГ 1. ОШИБКА: Путь не является директорией: %s", directory_path)
            return files

        # Выбираем метод обхода
        entries = dir_path.rglob("*") if recursive else dir_path.iterdir()

        for entry in entries:
            if entry.is_file():
                ext = entry.suffix.lower()
                if ext in ext_tuple:
                    files.append(str(entry.absolute()))
                    logger.debug("ШАГ 1. Найден файл: %s", entry.name)

        # Сортируем для детерминированности
        files = sorted(files)

        if not files:
            logger.warning("ШАГ 1. ПРЕДУПРЕЖДЕНИЕ: Не найдено файлов с расширениями %s", ext_tuple)
        else:
            logger.info("ШАГ 1. Получили %d файлов для индексации — УСПЕХ", len(files))

        return files


# ============================================================================
# 7. ASR-клиент вынесен в AI.llm_asr (реэкспорт см. секцию 16c)
# ============================================================================


# ============================================================================
# 8. Embeddings + Vector + Sparse вынесены в AI.llm_qdrant
# ============================================================================


# ============================================================================
# 8b. ingest_and_index вынесен в AI.llm_qdrant (реэкспорт см. секцию 16c)
# ============================================================================


# ============================================================================
# 9. Query Rewriters (UC-1: шаги 3.1, 3.2)
# ============================================================================


class RagQueryRewriter:
    """Генерирует N перефраз запроса для RAG через LLM (ШАГ 3.1)."""

    def __init__(self, llm_client: "OpenAIClient", count: int = 3) -> None:
        self._llm = llm_client
        self._count = count

    async def rewrite(self, query: str, ctx: RequestContext) -> List[str]:
        logger.info(
            "ШАГ 3.1. RagQueryRewriter — генерируем %d перефраз: request_id=%s",
            self._count, ctx.request_id,
        )
        prompt = (
            f"Перефразируй следующий поисковый запрос {self._count} разными способами. "
            f"Верни ТОЛЬКО JSON-массив строк, без объяснений.\n\nЗапрос: {query}"
        )
        request = LLMRequest(
            messages=[LLMMessage(role="user", content=prompt)],
            model=self._llm._default_model,
            temperature=0.7,
        )
        try:
            response = await self._llm.create_response(request, ctx)
            rephrases = json.loads(response.content)
            if isinstance(rephrases, list):
                logger.info(
                    "ШАГ 3.1. RagQueryRewriter — УСПЕХ: %d перефраз",
                    len(rephrases),
                )
                return [str(r) for r in rephrases[: self._count]]
        except Exception as exc:
            logger.warning(
                "ШАГ 3.1. RagQueryRewriter — ОШИБКА: %s, используем оригинал", exc,
            )
        return [query]


# ============================================================================
# 9b. CrawlerQueryRewriter вынесен в AI.llm_webcrawler (реэкспорт см. секцию 16c)
# ============================================================================


# ============================================================================
# 10. Retriever вынесен в AI.llm_qdrant (реэкспорт см. секцию 16c)
# ============================================================================


# ============================================================================
# 11. CrawlerClient вынесен в AI.llm_webcrawler (реэкспорт см. секцию 16c)
# ============================================================================


# ============================================================================
# 12. OpenAIClient (LLM-клиент)
# ============================================================================


class OpenAIClient:
    """Обертка над AsyncOpenAI: обычный вызов + стриминг + tool calling."""

    def __init__(
        self,
        api_key: str,
        base_url: Optional[str] = None,
        default_model: str = "gpt-4.1-mini",
    ) -> None:
        self._client = AsyncOpenAI(api_key=api_key, base_url=base_url)
        self._default_model = default_model

    async def create_response(
        self, request: LLMRequest, ctx: RequestContext,
    ) -> LLMResponse:
        model = request.model or self._default_model
        logger.info(
            "ШАГ LLM. Запрос к LLM (sync): request_id=%s, model=%s, messages=%d",
            ctx.request_id, model, len(request.messages),
        )

        kwargs: Dict[str, Any] = {
            "model": model,
            "messages": [self._msg_to_dict(m) for m in request.messages],
        }
        if request.temperature is not None:
            kwargs["temperature"] = request.temperature
        if request.max_output_tokens is not None:
            kwargs["max_tokens"] = request.max_output_tokens
        if request.tools:
            kwargs["tools"] = request.tools
            kwargs["tool_choice"] = request.tool_choice or "auto"
        if request.response_format:
            kwargs["response_format"] = request.response_format

        response = await self._client.chat.completions.create(**kwargs)
        choice = response.choices[0]
        # Cloud.ru иногда возвращает content=None при finish_reason=length, используем reasoning_content как запасной
        content = (choice.message.content
                   or getattr(choice.message, "reasoning_content", "")
                   or "")

        tool_calls: List[ToolCall] = []
        for tc in choice.message.tool_calls or []:
            tool_calls.append(ToolCall(
                id=str(tc.id),
                name=str(tc.function.name),
                arguments=json.loads(tc.function.arguments or "{}"),
            ))

        usage: Optional[TokenUsage] = None
        if hasattr(response, "usage") and response.usage:
            usage = TokenUsage(
                prompt_tokens=getattr(response.usage, "prompt_tokens", None),
                completion_tokens=getattr(response.usage, "completion_tokens", None),
                total_tokens=getattr(response.usage, "total_tokens", None),
            )

        logger.info(
            "ШАГ LLM. УСПЕХ: content_len=%d, tool_calls=%d, tokens=%s",
            len(content), len(tool_calls),
            usage.total_tokens if usage else "N/A",
        )

        raw = response.model_dump() if hasattr(response, "model_dump") else {}
        return LLMResponse(
            content=content, tool_calls=tool_calls,
            usage=usage, raw_response=raw,
        )

    async def stream_response(
        self, request: LLMRequest, ctx: RequestContext,
    ) -> AsyncIterator[StreamEvent]:
        model = request.model or self._default_model
        logger.info(
            "ШАГ LLM STREAM. Запрос к LLM (stream): request_id=%s, model=%s",
            ctx.request_id, model,
        )

        kwargs: Dict[str, Any] = {
            "model": model,
            "messages": [self._msg_to_dict(m) for m in request.messages],
            "stream": True,
        }
        if request.temperature is not None:
            kwargs["temperature"] = request.temperature
        if request.max_output_tokens is not None:
            kwargs["max_tokens"] = request.max_output_tokens
        if request.tools:
            kwargs["tools"] = request.tools
            kwargs["tool_choice"] = request.tool_choice or "auto"
        if request.response_format:
            kwargs["response_format"] = request.response_format

        collected_tool_calls: Dict[int, Dict[str, Any]] = {}

        stream = await self._client.chat.completions.create(**kwargs)
        async for chunk in stream:
            if not chunk.choices:
                continue
            choice = chunk.choices[0]
            delta = getattr(choice, "delta", None)

            if delta and getattr(delta, "content", None):
                yield StreamEvent(
                    kind="token_delta", data={"text": delta.content},
                )

            if delta and getattr(delta, "tool_calls", None):
                for tc_delta in delta.tool_calls:
                    idx = tc_delta.index
                    if idx not in collected_tool_calls:
                        collected_tool_calls[idx] = {
                            "id": "", "name": "", "arguments": "",
                        }
                    if tc_delta.id:
                        collected_tool_calls[idx]["id"] = tc_delta.id
                    if tc_delta.function:
                        if tc_delta.function.name:
                            collected_tool_calls[idx]["name"] = tc_delta.function.name
                        if tc_delta.function.arguments:
                            collected_tool_calls[idx]["arguments"] += (
                                tc_delta.function.arguments
                            )

        for idx in sorted(collected_tool_calls.keys()):
            tc_data = collected_tool_calls[idx]
            try:
                args = json.loads(tc_data["arguments"]) if tc_data["arguments"] else {}
            except json.JSONDecodeError:
                args = {}
            yield StreamEvent(kind="tool_call", data={
                "id": tc_data["id"],
                "name": tc_data["name"],
                "arguments": args,
            })

        logger.info(
            "ШАГ LLM STREAM. УСПЕХ: tool_calls=%d",
            len(collected_tool_calls),
        )

    @staticmethod
    def _msg_to_dict(msg: LLMMessage) -> Dict[str, Any]:
        d: Dict[str, Any] = {"role": msg.role, "content": msg.content}
        if msg.name:
            d["name"] = msg.name
        if msg.tool_call_id:
            d["tool_call_id"] = msg.tool_call_id
        if msg.tool_calls:
            d["tool_calls"] = msg.tool_calls
        return d


# ============================================================================
# 13. ChatOrchestrator — единая точка входа (UC-1 … UC-5)
# ============================================================================


class ChatOrchestrator:
    """Единый оркестратор для всех UC.

    Два публичных метода:
      - handle_user_input  — non-streaming (возвращает OrchestratorResult).
      - stream_user_input  — streaming (yield StreamEvent).
    """

    def __init__(
        self,
        context_builder: ContextBuilder,
        llm_client: OpenAIClient,
        tool_registry: ToolRegistry,
        tool_executor: ToolExecutor,
        retriever: Optional[Retriever] = None,
        rag_rewriter: Optional[RagQueryRewriter] = None,
        crawler_client: Optional[CrawlerClient] = None,
        crawler_rewriter: Optional[CrawlerQueryRewriter] = None,
        document_ingestor: Optional[DocumentIngestor] = None,
        asr_client: Optional[ASRClient] = None,
        json_repair: Optional[JsonRepairLLM] = None,
        strict_parser: Optional[StrictOutputParser] = None,
    ) -> None:
        self._builder = context_builder
        self._llm = llm_client
        self._tools = tool_registry
        self._tool_executor = tool_executor
        self._retriever = retriever
        self._rag_rewriter = rag_rewriter
        self._crawler = crawler_client
        self._crawler_rewriter = crawler_rewriter
        self._ingestor = document_ingestor
        self._asr = asr_client
        self._json_repair = json_repair or JsonRepairLLM()
        self._strict_parser = strict_parser or StrictOutputParser(self._json_repair)

    # ------------------------------------------------------------------
    # Non-streaming pipeline
    # ------------------------------------------------------------------

    async def handle_user_input(
        self, user_input: UserInput, ctx: RequestContext,
    ) -> OrchestratorResult:
        total_usage = TokenUsage()
        all_sources: List[Snippet] = []
        all_tool_calls: List[ToolCall] = []

        # --- ШАГ 0. Инициализация ---
        logger.info(
            "ШАГ 0. Инициализация запроса: request_id=%s, mode=%s, "
            "USE_INTERNET=%s, ENABLE_TOOLS=%s, ENABLE_ASR=%s",
            ctx.request_id, user_input.mode, user_input.use_internet,
            user_input.enable_tools, user_input.enable_asr,
        )

        ctx.mode = user_input.mode
        ctx.use_internet = user_input.use_internet
        ctx.enable_tools = user_input.enable_tools
        ctx.enable_streaming = user_input.enable_streaming
        ctx.enable_asr = user_input.enable_asr

        # ASR (UC-3)
        if user_input.enable_asr and user_input.audio_ref and self._asr:
            logger.info(
                "ШАГ 0.ASR. Распознаём аудио: request_id=%s, audio_ref=%s",
                ctx.request_id, user_input.audio_ref,
            )
            try:
                transcript = await self._asr.transcribe_file(
                    user_input.audio_ref, ctx,
                )
                if not transcript.text.strip():
                    logger.warning(
                        "ШАГ 0.ASR. Транскрипт пустой — аудио не распознано",
                    )
                    return OrchestratorResult(
                        response_text=(
                            "Аудио не распознано. Попробуйте загрузить файл ещё раз."
                        ),
                        context_used={"asr_empty": True},
                    )
                user_input.text = transcript.text
                logger.info(
                    "ШАГ 0.ASR. УСПЕХ: text_len=%d", len(transcript.text),
                )
            except Exception as exc:
                logger.error("ШАГ 0.ASR. ОШИБКА: %s", exc)
                return OrchestratorResult(
                    response_text="ASR недоступен. Попробуйте позже.",
                    context_used={"asr_error": str(exc)},
                )

        # Определяем conversation_id: оркестратор использует уже переданный ID
        ctx.conversation_id = user_input.conversation_id or ctx.conversation_id

        # --- ШАГ 2. Инжест документов (UC-2) ---
        doc_chunks: List[DocumentChunk] = []
        if user_input.files and self._ingestor:
            logger.info(
                "ШАГ 2. Инжест документов: request_id=%s, files=%d",
                ctx.request_id, len(user_input.files),
            )
            for file_ref in user_input.files:
                try:
                    chunks = await self._ingestor.ingest(file_ref, ctx)
                    doc_chunks.extend(chunks)
                except Exception as exc:
                    logger.error(
                        "ШАГ 2. ОШИБКА инжеста file_id=%s: %s",
                        file_ref.id, exc,
                    )
            logger.info("ШАГ 2. УСПЕХ: total_chunks=%d", len(doc_chunks))

        # --- ШАГ 3. RAG + Internet (UC-1, параллельно по флагам) ---
        rag_snippets: List[Snippet] = []
        internet_snippets: List[Snippet] = []
        is_rag_mode = user_input.mode in ("rag_tool", "rag_qa")
        has_query_text = bool(user_input.text)
        should_run_rag = bool(is_rag_mode and has_query_text and self._retriever)
        should_run_internet = bool(user_input.use_internet and has_query_text)

        logger.info(
            "ШАГ 3. Маршрутизация retrieval — request_id=%s mode=%s has_query_text=%s should_run_rag=%s should_run_internet=%s retriever_configured=%s crawler_configured=%s",
            ctx.request_id,
            user_input.mode,
            has_query_text,
            should_run_rag,
            should_run_internet,
            self._retriever is not None,
            self._crawler is not None,
        )

        async def _empty_snippets() -> List[Snippet]:
            return []

        rag_coro = self._run_rag(user_input.text, ctx) if should_run_rag else _empty_snippets()
        internet_coro = self._run_internet(user_input.text, ctx) if should_run_internet else _empty_snippets()

        if should_run_rag or should_run_internet:
            if should_run_rag and should_run_internet:
                logger.info(
                    "ШАГ 3. Параллельный запуск RAG + Internet — request_id=%s",
                    ctx.request_id,
                )
            elif should_run_rag:
                logger.info(
                    "ШАГ 3. Запуск только RAG — request_id=%s",
                    ctx.request_id,
                )
            else:
                logger.info(
                    "ШАГ 3. Запуск только Internet — request_id=%s",
                    ctx.request_id,
                )

            rag_result, internet_result = await asyncio.gather(
                rag_coro, internet_coro, return_exceptions=True,
            )

            if isinstance(rag_result, list):
                rag_snippets = rag_result
                all_sources.extend(rag_snippets)
            elif isinstance(rag_result, Exception):
                logger.warning("ШАГ 3. RAG ОШИБКА: %s", rag_result)

            if isinstance(internet_result, list):
                internet_snippets = internet_result
                all_sources.extend(internet_snippets)
            elif isinstance(internet_result, Exception):
                logger.warning("ШАГ 3. Internet ОШИБКА: %s", internet_result)

            logger.info(
                "ШАГ 3. Retrieval завершён — УСПЕХ: rag_snippets=%d, internet_snippets=%d",
                len(rag_snippets), len(internet_snippets),
            )
        else:
            logger.info(
                "ШАГ 3. Retrieval пропущен — request_id=%s reason=no_enabled_branches",
                ctx.request_id,
            )

        # --- ШАГ 4. Сбор контекста ---
        llm_request = await self._builder.build_input(
            user_input,
            rag_snippets or None,
            internet_snippets or None,
            doc_chunks or None,
            ctx,
            history=user_input.history,
        )

        if user_input.enable_tools:
            tools_for_model = self._tools.get_openai_tools(
                whitelist=ctx.tool_whitelist,
            )
            if tools_for_model:
                llm_request.tools = tools_for_model
                llm_request.tool_choice = "auto"

        # --- ШАГ 5. LLM + Tool calling цикл (UC-5) ---
        logger.info("ШАГ 5. Первый вызов LLM: request_id=%s", ctx.request_id)
        response = await self._llm.create_response(llm_request, ctx)
        self._accumulate_usage(total_usage, response.usage)

        tool_iteration = 0
        while response.tool_calls and tool_iteration < ctx.tool_call_limit:
            tool_iteration += 1
            logger.info(
                "ШАГ 5.%d. Tool calls: count=%d, iteration=%d/%d",
                tool_iteration, len(response.tool_calls),
                tool_iteration, ctx.tool_call_limit,
            )

            tool_results = await self._tool_executor.execute_many(
                response.tool_calls, ctx,
            )
            all_tool_calls.extend(response.tool_calls)

            assistant_tc_msg = LLMMessage(
                role="assistant",
                content=response.content or "",
                tool_calls=[
                    {
                        "id": tc.id, "type": "function",
                        "function": {
                            "name": tc.name,
                            "arguments": json.dumps(
                                tc.arguments, ensure_ascii=False,
                            ),
                        },
                    }
                    for tc in response.tool_calls
                ],
            )
            llm_request.messages.append(assistant_tc_msg)

            for tr in tool_results:
                content_str = (
                    json.dumps(tr.output, ensure_ascii=False)
                    if not tr.is_error
                    else json.dumps({"error": tr.error_message}, ensure_ascii=False)
                )
                llm_request.messages.append(LLMMessage(
                    role="tool", content=content_str, tool_call_id=tr.id,
                ))
                if tr.is_error:
                    logger.warning(
                        "ШАГ 5.%d. Tool %s ошибка: %s",
                        tool_iteration, tr.name, tr.error_message,
                    )

            logger.info(
                "ШАГ 5.%d. Повторный вызов LLM: request_id=%s",
                tool_iteration, ctx.request_id,
            )
            response = await self._llm.create_response(llm_request, ctx)
            self._accumulate_usage(total_usage, response.usage)

        if response.tool_calls and tool_iteration >= ctx.tool_call_limit:
            logger.warning(
                "ШАГ 5. Лимит tool calls (%d) — частичный результат",
                ctx.tool_call_limit,
            )

        final_text = response.content

        # --- ШАГ 6. JSON repair (если нужно) ---
        if "output_schema" in llm_request.metadata:
            schema: type[BaseModel] = llm_request.metadata["output_schema"]
            try:
                await self._strict_parser.parse_json(final_text, schema, ctx)
            except Exception as exc:
                logger.warning(
                    "ШАГ 6. Валидация JSON-ответа не прошла: %s", exc,
                )

        return OrchestratorResult(
            response_text=final_text,
            sources=all_sources,
            tool_calls=all_tool_calls,
            tokens_used=total_usage if total_usage.total_tokens else None,
            context_used={
                "mode": user_input.mode,
                "used_rag": bool(rag_snippets),
                "used_internet": bool(internet_snippets),
                "used_tools": bool(all_tool_calls),
                "used_documents": bool(doc_chunks),
            },
        )

    # ------------------------------------------------------------------
    # Streaming pipeline (UC-4)
    # ------------------------------------------------------------------

    async def stream_user_input(
        self, user_input: UserInput, ctx: RequestContext,
    ) -> AsyncIterator[StreamEvent]:
        yield StreamEvent(
            kind="init", step="init",
            data={"request_id": ctx.request_id},
        )

        ctx.mode = user_input.mode
        ctx.use_internet = user_input.use_internet
        ctx.enable_tools = user_input.enable_tools
        ctx.enable_streaming = True
        ctx.enable_asr = user_input.enable_asr

        # ШАГ 0. ASR
        if user_input.enable_asr and user_input.audio_ref and self._asr:
            try:
                transcript = await self._asr.transcribe_file(
                    user_input.audio_ref, ctx,
                )
                if not transcript.text.strip():
                    yield StreamEvent(
                        kind="error",
                        data={"message": "Аудио не распознано"},
                    )
                    yield StreamEvent(kind="done", step="done", data={})
                    return
                user_input.text = transcript.text
            except Exception as exc:
                yield StreamEvent(
                    kind="error",
                    data={"message": f"ASR ошибка: {exc}"},
                )
                yield StreamEvent(kind="done", step="done", data={})
                return

        # Определяем conversation_id: оркестратор использует уже переданный ID
        ctx.conversation_id = user_input.conversation_id or ctx.conversation_id

        # ШАГ 2. Инжест документов
        doc_chunks: List[DocumentChunk] = []
        if user_input.files and self._ingestor:
            for file_ref in user_input.files:
                try:
                    chunks = await self._ingestor.ingest(file_ref, ctx)
                    doc_chunks.extend(chunks)
                except Exception as exc:
                    logger.error(
                        "ШАГ (stream) 2. ОШИБКА инжеста: %s", exc,
                    )

        # ШАГ 3. RAG + Internet
        rag_snippets: List[Snippet] = []
        internet_snippets: List[Snippet] = []
        is_rag_mode = user_input.mode in ("rag_tool", "rag_qa")
        has_query_text = bool(user_input.text)
        should_run_rag = bool(is_rag_mode and has_query_text and self._retriever)
        should_run_internet = bool(user_input.use_internet and has_query_text)

        logger.info(
            "ШАГ (stream) 3. Маршрутизация retrieval — request_id=%s mode=%s has_query_text=%s should_run_rag=%s should_run_internet=%s retriever_configured=%s crawler_configured=%s",
            ctx.request_id,
            user_input.mode,
            has_query_text,
            should_run_rag,
            should_run_internet,
            self._retriever is not None,
            self._crawler is not None,
        )

        if should_run_rag:
            yield StreamEvent(
                kind="rag_progress", step="rag_start", data={},
            )
            try:
                rag_snippets = await self._run_rag(user_input.text, ctx)
            except Exception as exc:
                logger.warning("ШАГ (stream) 3. RAG ОШИБКА: %s", exc)

        if should_run_internet:
            yield StreamEvent(
                kind="internet_progress", step="internet_start", data={},
            )
            try:
                internet_snippets = await self._run_internet(
                    user_input.text, ctx,
                )
            except Exception as exc:
                logger.warning(
                    "ШАГ (stream) 3. Internet ОШИБКА: %s", exc,
                )

        if not should_run_rag and not should_run_internet:
            logger.info(
                "ШАГ (stream) 3. Retrieval пропущен — request_id=%s reason=no_enabled_branches",
                ctx.request_id,
            )

        # ШАГ 4. Context
        llm_request = await self._builder.build_input(
            user_input,
            rag_snippets or None,
            internet_snippets or None,
            doc_chunks or None,
            ctx,
            history=user_input.history,
        )

        if user_input.enable_tools:
            tools_for_model = self._tools.get_openai_tools(
                whitelist=ctx.tool_whitelist,
            )
            if tools_for_model:
                llm_request.tools = tools_for_model
                llm_request.tool_choice = "auto"

        # ШАГ 5-6. Stream LLM
        collected_text = ""
        tool_calls_from_stream: List[ToolCall] = []

        async for event in self._llm.stream_response(llm_request, ctx):
            if event.kind == "token_delta":
                collected_text += event.data.get("text", "")
                yield event
            elif event.kind == "tool_call":
                tc = ToolCall(
                    id=event.data.get("id", str(uuid.uuid4())),
                    name=event.data.get("name", ""),
                    arguments=event.data.get("arguments", {}),
                )
                tool_calls_from_stream.append(tc)
                yield event

        # Tool results + re-call
        if tool_calls_from_stream and user_input.enable_tools:
            tool_results = await self._tool_executor.execute_many(
                tool_calls_from_stream, ctx,
            )
            for tr in tool_results:
                yield StreamEvent(
                    kind="tool_result",
                    data={
                        "name": tr.name, "output": tr.output,
                        "is_error": tr.is_error,
                    },
                )

            assistant_tc_msg = LLMMessage(
                role="assistant", content=collected_text,
                tool_calls=[
                    {
                        "id": tc.id, "type": "function",
                        "function": {
                            "name": tc.name,
                            "arguments": json.dumps(
                                tc.arguments, ensure_ascii=False,
                            ),
                        },
                    }
                    for tc in tool_calls_from_stream
                ],
            )
            llm_request.messages.append(assistant_tc_msg)

            for tr in tool_results:
                content_str = (
                    json.dumps(tr.output, ensure_ascii=False)
                    if not tr.is_error
                    else json.dumps({"error": tr.error_message})
                )
                llm_request.messages.append(LLMMessage(
                    role="tool", content=content_str, tool_call_id=tr.id,
                ))

            collected_text = ""
            async for event in self._llm.stream_response(llm_request, ctx):
                if event.kind == "token_delta":
                    collected_text += event.data.get("text", "")
                    yield event

        yield StreamEvent(
            kind="final", step="final",
            data={
                "text": collected_text,
                "sources": [s.model_dump() for s in (rag_snippets + internet_snippets)],
                "tool_calls": [tc.model_dump() for tc in tool_calls_from_stream],
                "history_messages": len(user_input.history) if user_input.history else 0,
            },
        )
        yield StreamEvent(
            kind="done", step="done",
            data={
                "conversation_id": ctx.conversation_id,
                "sources": [s.model_dump() for s in (rag_snippets + internet_snippets)],
                "tool_calls": [tc.model_dump() for tc in tool_calls_from_stream],
            },
        )

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    async def _run_rag(
        self, query: str, ctx: RequestContext,
    ) -> List[Snippet]:
        logger.info(
            "ШАГ 3.1-3.3. RAG pipeline: request_id=%s", ctx.request_id,
        )
        queries = [query]
        if self._rag_rewriter:
            try:
                queries = await self._rag_rewriter.rewrite(query, ctx)
            except Exception as exc:
                logger.warning(
                    "ШАГ 3.1. QueryRewriter ОШИБКА: %s — используем оригинал", exc,
                )
                queries = [query]

        snippets = await self._retriever.retrieve_multi(  # type: ignore[union-attr]
            queries, ctx, top_k=ctx.rag_top_k,
            filters=ctx.rag_filters,
        )
        logger.info("ШАГ 3.3. RAG — УСПЕХ: %d сниппетов", len(snippets))
        return snippets

    async def _run_internet(
        self, query: str, ctx: RequestContext,
    ) -> List[Snippet]:
        if not self._crawler:
            logger.info(
                "ШАГ 3.4. Internet — пропускаем (CrawlerClient не настроен)",
            )
            return []

        logger.info(
            "ШАГ 3.2-3.4. Internet pipeline: request_id=%s", ctx.request_id,
        )
        queries = [query]
        if self._crawler_rewriter:
            try:
                queries = await self._crawler_rewriter.rewrite(query, ctx)
            except Exception as exc:
                logger.warning(
                    "ШАГ 3.2. CrawlerQueryRewriter ОШИБКА: %s", exc,
                )
                queries = [query]

        snippets = await self._crawler.search(queries, ctx)
        logger.info(
            "ШАГ 3.4. Internet — УСПЕХ: %d сниппетов", len(snippets),
        )
        return snippets

    @staticmethod
    def _accumulate_usage(
        total: TokenUsage, usage: Optional[TokenUsage],
    ) -> None:
        if not usage:
            return
        if usage.prompt_tokens:
            total.prompt_tokens = (total.prompt_tokens or 0) + usage.prompt_tokens
        if usage.completion_tokens:
            total.completion_tokens = (
                (total.completion_tokens or 0) + usage.completion_tokens
            )
        if usage.total_tokens:
            total.total_tokens = (total.total_tokens or 0) + usage.total_tokens


# ============================================================================
# 14. Фабричные функции (создание из переменных окружения CLOUDRU_*)
# ============================================================================


def create_cloudru_openai_client_from_env() -> OpenAIClient:
    """OpenAIClient → Cloud.ru (GLM-4.7)."""
    api_key = os.environ["CLOUDRU_API_KEY"]
    base_url = os.environ.get(
        "CLOUDRU_BASE_URL", "https://foundation-models.api.cloud.ru/v1",
    )
    model_name = os.environ.get("CLOUDRU_MODEL_NAME", "zai-org/GLM-4.7")
    logger.info(
        "Создаём OpenAIClient: base_url=%s, model=%s", base_url, model_name,
    )
    return OpenAIClient(api_key=api_key, base_url=base_url, default_model=model_name)


def create_cloudru_asr_client_from_env() -> ASRClient:
    """ASRClient → Cloud.ru whisper."""
    api_key = os.environ["CLOUDRU_API_KEY"]
    base_url = os.environ.get(
        "CLOUDRU_BASE_URL", "https://foundation-models.api.cloud.ru/v1",
    )
    model = os.environ.get("ASR_MODEL", "openai/whisper-large-v3")
    asr_language = os.environ.get("ASR_LANGUAGE", "ru")
    asr_timeout_seconds = float(os.environ.get("ASR_HTTP_TIMEOUT_SECONDS", "180"))
    asr_max_payload_bytes = int(
        os.environ.get(
            "ASR_MAX_PAYLOAD_BYTES",
            str(ASRClient.DEFAULT_SAFE_UPLOAD_BYTES),
        ),
    )
    asr_chunk_duration_seconds = int(
        os.environ.get(
            "ASR_CHUNK_DURATION_SECONDS",
            str(ASRClient.DEFAULT_CHUNK_DURATION_SECONDS),
        ),
    )
    asr_target_bitrate_kbps = int(
        os.environ.get("ASR_TARGET_BITRATE_KBPS", "48"),
    )
    logger.info(
        "Создаём ASRClient: base_url=%s, model=%s, language=%s, timeout=%.1fs, "
        "max_payload=%d, chunk_duration=%ds, bitrate=%dkbps",
        base_url,
        model,
        asr_language,
        asr_timeout_seconds,
        asr_max_payload_bytes,
        asr_chunk_duration_seconds,
        asr_target_bitrate_kbps,
    )
    return ASRClient(
        api_key=api_key,
        base_url=base_url,
        model=model,
        language=asr_language,
        timeout_seconds=asr_timeout_seconds,
        max_payload_bytes=asr_max_payload_bytes,
        chunk_duration_seconds=asr_chunk_duration_seconds,
        target_bitrate_kbps=asr_target_bitrate_kbps,
    )


def create_cloudru_embedding_client_from_env() -> CloudRuEmbeddingClient:
    """CloudRuEmbeddingClient с моделью из CLOUDRU_EMBED_MODEL и API-ключом."""
    api_key = os.environ["CLOUDRU_API_KEY"]
    base_url = os.environ.get(
        "CLOUDRU_BASE_URL", "https://foundation-models.api.cloud.ru/v1",
    )
    embed_model = os.environ.get(
        "CLOUDRU_EMBED_MODEL", "Qwen/Qwen3-Embedding-0.6B",
    )
    logger.info(
        "Создаём CloudRuEmbeddingClient: base_url=%s, model=%s",
        base_url, embed_model,
    )
    return CloudRuEmbeddingClient(
        api_key=api_key, base_url=base_url, model_name=embed_model,
    )


def create_sparse_embedding_client_from_env() -> SparseEmbeddingClient:
    """SparseEmbeddingClient из SPARSE_EMBED_MODEL (по умолчанию opensearch multilingual)."""
    model_name = os.environ.get(
        "SPARSE_EMBED_MODEL",
        "opensearch-project/opensearch-neural-sparse-encoding-multilingual-v1",
    )
    logger.info("Создаём SparseEmbeddingClient: model=%s", model_name)
    return SparseEmbeddingClient(model_name=model_name)


def create_default_orchestrator_from_env(
    collection_name: str = "default",
    crawler_base_url: Optional[str] = None,
    enable_sparse: bool = True,
) -> ChatOrchestrator:
    """Полная сборка ChatOrchestrator из переменных окружения CLOUDRU_*.

    Параметры:
      - crawler_base_url: URL crawler-сервера. Если None — берётся из CRAWLER_BASE_URL.
      - enable_sparse: если True — создаёт SparseEmbeddingClient и передаёт
        в Retriever для hybrid (dense + sparse) search.
    """
    llm_client = create_cloudru_openai_client_from_env()
    model_name = os.environ.get("CLOUDRU_MODEL_NAME", "zai-org/GLM-4.7")

    embedding_client = create_cloudru_embedding_client_from_env()
    sparse_client: Optional[SparseEmbeddingClient] = None
    if enable_sparse:
        sparse_client = create_sparse_embedding_client_from_env()

    vector_store = QdrantVectorStore(
        host=os.environ.get("QDRANT_HOST", "localhost"),
        port=int(os.environ.get("QDRANT_PORT", "6334")),
    )
    retriever = Retriever(
        embedding_client, vector_store, collection=collection_name,
        sparse_client=sparse_client,
    )

    context_builder = ContextBuilder(default_model=model_name)
    tool_registry = ToolRegistry()
    tool_executor = ToolExecutor(tool_registry)

    rag_rewriter = RagQueryRewriter(llm_client, count=3)
    crawler_rewriter = CrawlerQueryRewriter(llm_client, count=2)
    resolved_crawler_url = crawler_base_url or os.environ.get("CRAWLER_BASE_URL")
    crawler = CrawlerClient(base_url=resolved_crawler_url) if resolved_crawler_url else None
    logger.info(
        "CrawlerClient: base_url=%s (configured=%s)",
        resolved_crawler_url, bool(crawler),
    )

    asr_client = create_cloudru_asr_client_from_env()
    ingestor = DocumentIngestor()
    json_repair = JsonRepairLLM()
    strict_parser = StrictOutputParser(json_repair)

    logger.info(
        "Создаём ChatOrchestrator: model=%s, collection=%s, sparse=%s",
        model_name, collection_name, bool(sparse_client),
    )

    return ChatOrchestrator(
        context_builder=context_builder,
        llm_client=llm_client,
        tool_registry=tool_registry,
        tool_executor=tool_executor,
        retriever=retriever,
        rag_rewriter=rag_rewriter,
        crawler_client=crawler,
        crawler_rewriter=crawler_rewriter,
        document_ingestor=ingestor,
        asr_client=asr_client,
        json_repair=json_repair,
        strict_parser=strict_parser,
    )


# ============================================================================
# 15. Утилиты для UC (отчеты, вспомогательные функции)
# ============================================================================


def generate_indexing_report(
    files_processed: int,
    total_chunks: int,
    collection: str,
    status: str,
    error: Optional[str] = None,
    qdrant_host: str = "localhost",
    qdrant_port: int = 6334,
) -> str:
    """ШАГ 6. Формирование JSON-отчета с результатами индексации.

    Возвращает структурированный ответ с полями:
    - processed: количество обработанных файлов
    - failed: количество файлов с ошибками
    - total_chunks: общее количество чанков
    - collection: имя коллекции
    - qdrant: адрес хранилища
    - status: success/error

    Args:
        files_processed: Количество обработанных файлов
        total_chunks: Общее количество чанков
        collection: Имя коллекции Qdrant
        status: Статус индексации ("success" или "error")
        error: Сообщение об ошибке (если status != "success")
        qdrant_host: Хост Qdrant
        qdrant_port: Порт Qdrant

    Returns:
        JSON-строка с отчётом

    Логирование:
        «ШАГ 6. Пакетная индексация завершена — УСПЕХ/ОШИБКАМИ»
    """
    logger.info("ШАГ 6. Формирование отчета")

    report: Dict[str, Any] = {
        "processed": files_processed,
        "failed": 0 if status == "success" else files_processed,
        "total_chunks": total_chunks,
        "collection": collection,
        "qdrant": f"{qdrant_host}:{qdrant_port}",
        "status": status,
    }

    if status != "success" and error:
        report["error"] = error
        logger.error("ШАГ 6. Пакетная индексация завершена с ОШИБКАМИ")
    else:
        logger.info("ШАГ 6. Пакетная индексация завершена — УСПЕХ")

    return json.dumps(report, ensure_ascii=False, indent=2)


# ============================================================================
# 16. UC-хелперы (загрузка конфига, промпт-шаблоны, LLM-обёртки)
# ============================================================================


def load_env_and_validate(
    env_dir: str,
    required_keys: Optional[List[str]] = None,
) -> Dict[str, str]:
    """Загрузка .env из указанной директории + валидация обязательных ключей.

    Возвращает словарь со ВСЕМИ стандартными Cloud.ru переменными
    (с дефолтами). Если обязательный ключ отсутствует — бросает ValueError.

    Используется во ВСЕХ UC-скриптах вместо повторяющегося блока
    load_dotenv + os.environ.get + проверка.

    Args:
        env_dir: Путь к директории с файлом .env
        required_keys: Список обязательных ключей (по умолчанию ["CLOUDRU_API_KEY"])

    Returns:
        Dict со всеми значениями переменных окружения

    Raises:
        ValueError: если обязательный ключ не найден
    """
    logger.info("ШАГ ENV. Загрузка .env из %s", env_dir)

    try:
        from dotenv import load_dotenv
        env_path = Path(env_dir) / ".env"
        if env_path.exists():
            load_dotenv(env_path)
            logger.info("ШАГ ENV. Загружен .env из %s — УСПЕХ", env_path)
        else:
            logger.warning("ШАГ ENV. Файл .env не найден в %s, используем переменные окружения", env_path)
    except ImportError:
        logger.warning("ШАГ ENV. python-dotenv не установлен, используем переменные окружения")

    _required = required_keys or ["CLOUDRU_API_KEY"]
    for key in _required:
        if not os.environ.get(key):
            msg = f"Обязательная переменная {key} не задана"
            logger.error("ШАГ ENV. ОШИБКА: %s", msg)
            raise ValueError(msg)

    config = {
        "CLOUDRU_API_KEY": os.environ["CLOUDRU_API_KEY"],
        "CLOUDRU_BASE_URL": os.environ.get(
            "CLOUDRU_BASE_URL", "https://foundation-models.api.cloud.ru/v1",
        ),
        "CLOUDRU_MODEL_NAME": os.environ.get("CLOUDRU_MODEL_NAME", "zai-org/GLM-4.7"),
        "CLOUDRU_EMBED_MODEL": os.environ.get("CLOUDRU_EMBED_MODEL", "Qwen/Qwen3-Embedding-0.6B"),
        "QDRANT_HOST": os.environ.get("QDRANT_HOST", "localhost"),
        "QDRANT_PORT": os.environ.get("QDRANT_PORT", "6334"),
        "SPARSE_CACHE_DIR": os.environ.get("SPARSE_CACHE_DIR", ""),
        "CRAWLER_BASE_URL": os.environ.get("CRAWLER_BASE_URL", ""),
        "ASR_MODEL": os.environ.get("ASR_MODEL", "openai/whisper-large-v3"),
        "ASR_LANGUAGE": os.environ.get("ASR_LANGUAGE", "ru"),
        "ASR_HTTP_TIMEOUT_SECONDS": os.environ.get("ASR_HTTP_TIMEOUT_SECONDS", "180"),
        "ASR_MAX_PAYLOAD_BYTES": os.environ.get(
            "ASR_MAX_PAYLOAD_BYTES",
            str(ASRClient.DEFAULT_SAFE_UPLOAD_BYTES),
        ),
        "ASR_CHUNK_DURATION_SECONDS": os.environ.get("ASR_CHUNK_DURATION_SECONDS", "300"),
        "ASR_TARGET_BITRATE_KBPS": os.environ.get("ASR_TARGET_BITRATE_KBPS", "48"),
    }

    logger.info("ШАГ ENV. Конфигурация загружена — УСПЕХ")
    for k, v in config.items():
        if k != "CLOUDRU_API_KEY":
            logger.info("  %s = %s", k, v or "(не задано)")

    return config


def create_rag_clients_from_env(
    collection: str = "default",
    sparse_cache_dir: Optional[str] = None,
) -> Dict[str, Any]:
    """Фабрика RAG-клиентов из переменных окружения.

    Создаёт и возвращает словарь с ключами:
      - "embedding": CloudRuEmbeddingClient
      - "sparse": SparseEmbeddingClient
      - "vector_store": QdrantVectorStore
      - "retriever": Retriever (dense + sparse)

    Args:
        collection: Имя коллекции Qdrant
        sparse_cache_dir: Путь к кэшу sparse-модели (если None — из SPARSE_CACHE_DIR)

    Returns:
        Dict с клиентами

    Raises:
        ValueError: при отсутствии CLOUDRU_API_KEY
    """
    logger.info("ШАГ RAG INIT. Создание RAG-клиентов для коллекции '%s'", collection)

    embedding = create_cloudru_embedding_client_from_env()

    _cache = sparse_cache_dir or os.environ.get("SPARSE_CACHE_DIR", "")
    sparse = SparseEmbeddingClient(cache_dir=_cache) if _cache else None

    vector_store = QdrantVectorStore(
        host=os.environ.get("QDRANT_HOST", "localhost"),
        port=int(os.environ.get("QDRANT_PORT", "6334")),
    )

    retriever = Retriever(
        embedding_client=embedding,
        vector_store=vector_store,
        collection=collection,
        sparse_client=sparse,
    )

    clients = {
        "embedding": embedding,
        "sparse": sparse,
        "vector_store": vector_store,
        "retriever": retriever,
    }

    logger.info(
        "ШАГ RAG INIT. УСПЕХ: embedding=%s, sparse=%s, collection=%s",
        type(embedding).__name__, bool(sparse), collection,
    )
    return clients


def build_rag_prompt(
    query: str,
    snippets: List[Dict[str, Any]],
) -> str:
    """Формирует промпт для LLM на основе RAG-сниппетов.

    Общий шаблон для UC-2 и любых других RAG-скриптов.

    Args:
        query: Вопрос пользователя
        snippets: Список словарей с ключами text, source_id, score

    Returns:
        Готовый промпт-строка
    """
    logger.info("ШАГ PROMPT (RAG). Сборка промпта из %d сниппетов", len(snippets))

    context_parts = []
    for i, s in enumerate(snippets, 1):
        source = s.get("source_id", "unknown")
        text = s.get("text", "").strip()
        score = s.get("score", 0.0)
        if text:
            context_parts.append(
                f"=== ИСТОЧНИК {i} [релевантность: {score:.3f}] ===\n"
                f"ID: {source}\nТекст: {text}\n"
            )

    context_text = (
        "\n\n".join(context_parts)
        if context_parts
        else "[Нет релевантных документов в базе знаний]"
    )

    prompt = (
        "Ты — интеллектуальный ассистент для анализа научных документов и данных.\n"
        "Твоя задача — дать точный, обоснованный ответ на вопрос пользователя, "
        "используя ТОЛЬКО предоставленный контекст из базы знаний.\n\n"
        f"═══ КОНТЕКСТ ИЗ БАЗЫ ЗНАНИЙ (найдено {len(snippets)} фрагментов) ═══\n\n"
        f"{context_text}\n\n"
        f"═══ ВОПРОС ПОЛЬЗОВАТЕЛЯ ═══\n\n{query}\n\n"
        "═══ ИНСТРУКЦИИ ═══\n"
        "1. Отвечай ТОЛЬКО на основе предоставленных фрагментов.\n"
        "2. Если информации недостаточно — честно сообщи об этом.\n"
        "3. Указывай источники: [Источник N].\n"
        "4. Структурируй: краткая сводка → подробный разбор → выводы.\n"
        "5. Сохраняй научную точность и терминологию из контекста.\n"
        "6. Отвечай на языке вопроса.\n"
    )

    logger.info("ШАГ PROMPT (RAG). УСПЕХ: prompt_len=%d", len(prompt))
    return prompt


def build_web_prompt(
    query: str,
    snippets: List[Any],
    chunks: List[Any],
) -> str:
    """Формирует промпт для LLM на основе веб-контента (UC-3).

    Args:
        query: Исходный поисковый запрос
        snippets: Список Snippet (с .source_id, .metadata)
        chunks: Список DocumentChunk (с .text, .source_id)

    Returns:
        Готовый промпт-строка
    """
    logger.info(
        "ШАГ PROMPT (WEB). Сборка промпта: %d сниппетов, %d чанков",
        len(snippets), len(chunks),
    )

    source_parts = []
    for i, s in enumerate(snippets, 1):
        url = getattr(s, "source_id", None) or "unknown"
        title = (getattr(s, "metadata", None) or {}).get("title", "")
        line = f"[{i}] {url}"
        if title:
            line += f" — {title[:80]}"
        source_parts.append(line)

    sources_text = "\n".join(source_parts) if source_parts else "[Источники не найдены]"

    content_parts = []
    for i, chunk in enumerate(chunks[:20], 1):
        src = getattr(chunk, "source_id", None) or "unknown"
        text = getattr(chunk, "text", "").strip()
        if text:
            content_parts.append(
                f"=== ФРАГМЕНТ {i} [источник: {src}] ===\n{text}\n"
            )

    content_text = (
        "\n\n".join(content_parts)
        if content_parts
        else "[Контент не извлечён]"
    )

    prompt = (
        "Ты — интеллектуальный ассистент для поиска и анализа информации из интернета.\n"
        "Твоя задача — дать точный, обоснованный ответ, используя ТОЛЬКО предоставленный "
        "контент из веб-источников.\n\n"
        f"═══ ИСТОЧНИКИ (найдено {len(snippets)} страниц) ═══\n\n{sources_text}\n\n"
        f"═══ КОНТЕКСТ ИЗ ВЕБ-ИСТОЧНИКОВ ({len(chunks)} фрагментов) ═══\n\n{content_text}\n\n"
        f"═══ ВОПРОС ПОЛЬЗОВАТЕЛЯ ═══\n\n{query}\n\n"
        "═══ ИНСТРУКЦИИ ═══\n"
        "1. Отвечай ТОЛЬКО на основе предоставленных фрагментов из интернета.\n"
        "2. Если информации недостаточно — честно сообщи об этом.\n"
        "3. Указывай источники: [N].\n"
        "4. Структурируй: краткая сводка → подробный разбор → выводы.\n"
        "5. Укажи даты если они есть в контексте.\n"
        "6. Отвечай на языке вопроса.\n"
    )

    logger.info("ШАГ PROMPT (WEB). УСПЕХ: prompt_len=%d", len(prompt))
    return prompt


async def stream_llm_to_stdout(
    llm_client: "OpenAIClient",
    prompt: str,
    ctx: "RequestContext",
    system_message: str = "Ты — полезный ассистент.",
    model: Optional[str] = None,
    temperature: float = 0.3,
    max_output_tokens: int = 4096,
) -> Optional[str]:
    """Стриминг ответа LLM с выводом токенов в stdout.

    Общий хелпер для UC-2, UC-3, UC-4 и любых других скриптов,
    которым нужен «живой» вывод ответа.

    Args:
        llm_client: OpenAIClient
        prompt: Промпт (user message)
        ctx: RequestContext
        system_message: Системное сообщение
        model: Модель (если None — дефолт клиента)
        temperature: Температура
        max_output_tokens: Лимит токенов

    Returns:
        Собранный текст ответа или None при ошибке
    """
    logger.info(
        "ШАГ LLM STREAM (stdout). model=%s, prompt_len=%d, request_id=%s",
        model or "(default)", len(prompt), ctx.request_id,
    )

    try:
        request = LLMRequest(
            messages=[
                LLMMessage(role="system", content=system_message),
                LLMMessage(role="user", content=prompt),
            ],
            model=model or "",
            temperature=temperature,
            max_output_tokens=max_output_tokens,
        )

        print("\n" + "=" * 70)
        print("LLM ОТВЕЧАЕТ (стриминг):")
        print("=" * 70)

        collected: List[str] = []
        token_count = 0

        async for event in llm_client.stream_response(request, ctx):
            if event.kind == "token_delta":
                text = event.data.get("text", "")
                collected.append(text)
                token_count += 1
                print(text, end="", flush=True)

        print("\n" + "=" * 70)

        answer = "".join(collected).strip()
        if not answer:
            logger.warning("ШАГ LLM STREAM (stdout). LLM вернул пустой ответ")
            return None

        logger.info(
            "ШАГ LLM STREAM (stdout). УСПЕХ: tokens=%d, answer_len=%d",
            token_count, len(answer),
        )
        return answer

    except Exception as exc:
        logger.error("ШАГ LLM STREAM (stdout). ОШИБКА: %s", exc)
        return None


async def query_llm_simple(
    llm_client: "OpenAIClient",
    prompt: str,
    ctx: "RequestContext",
    system_message: str = "Ты — полезный ассистент.",
    model: Optional[str] = None,
    temperature: float = 0.3,
    max_output_tokens: int = 4096,
) -> Optional[str]:
    """Non-streaming запрос к LLM, возвращает текст ответа.

    Общий хелпер для UC-скриптов, где не нужен стриминг.

    Args:
        llm_client: OpenAIClient
        prompt: Промпт (user message)
        ctx: RequestContext
        system_message: Системное сообщение
        model: Модель (если None — дефолт клиента)
        temperature: Температура
        max_output_tokens: Лимит токенов

    Returns:
        Текст ответа или None при ошибке
    """
    logger.info(
        "ШАГ LLM SIMPLE. model=%s, prompt_len=%d, request_id=%s",
        model or "(default)", len(prompt), ctx.request_id,
    )

    try:
        request = LLMRequest(
            messages=[
                LLMMessage(role="system", content=system_message),
                LLMMessage(role="user", content=prompt),
            ],
            model=model or "",
            temperature=temperature,
            max_output_tokens=max_output_tokens,
        )

        response = await llm_client.create_response(request, ctx)
        answer = response.content.strip()

        if not answer:
            logger.warning("ШАГ LLM SIMPLE. LLM вернул пустой ответ")
            return None

        logger.info("ШАГ LLM SIMPLE. УСПЕХ: answer_len=%d", len(answer))
        return answer

    except Exception as exc:
        logger.error("ШАГ LLM SIMPLE. ОШИБКА: %s", exc)
        return None


# ============================================================================
# 16b. Semantic Scholar вынесен в AI.llm_semantic_scholar
# ============================================================================


# source-of-truth: AI.llm_semantic_scholar (см. реэкспорт в секции 16c)


# ============================================================================
# 16c. Re-export вынесенных подсистем (фасадный слой, source-of-truth в модулях)
# ============================================================================

from AI.llm_asr import ASRClient, Transcript, TranscriptSegment
from AI.llm_qdrant import (
    CloudRuEmbeddingClient,
    EmbeddingClient,
    HybridSearchConfig,
    QdrantCollectionConfig,
    QdrantVectorStore,
    Retriever,
    SparseEmbeddingClient,
    SparseVectorData,
    VectorStore,
    ingest_and_index,
)
from AI.llm_semantic_scholar import (
    S2_AUTHOR_DETAIL_FIELDS,
    S2_AUTHOR_SEARCH_FIELDS,
    S2_CITATION_FIELDS,
    S2_PAPER_FIELDS,
    S2_RATE_LIMIT_DELAY,
    S2_VALID_FIELDS_OF_STUDY,
    S2Client,
    S2FieldInference,
    S2SearchFilter,
)
from AI.llm_webcrawler import CrawlerClient, CrawlerQueryRewriter


# ============================================================================
# 17. __all__
# ============================================================================

__all__ = [
    # DTO
    "RequestContext", "UserInput", "LLMMessage", "LLMRequest", "LLMResponse",
    "TokenUsage", "StreamEvent", "ToolCall", "ToolResult",
    "ChatMessage", "DocumentChunk", "IndexMetadata", "Snippet",
    "Transcript", "TranscriptSegment", "FileRef", "OrchestratorResult",
    # Qdrant DTO
    "SparseVectorData", "QdrantCollectionConfig", "HybridSearchConfig",
    # Context
    "ContextBuilder",
    # JSON repair + validation
    "JsonRepairLLM", "StrictOutputParser",
    # Tools
    "ToolSpec", "ToolRequestContext", "ToolRegistry", "ToolExecutor",
    # Ingestion
    "Chunker", "DocumentIngestor",
    # ASR
    "ASRClient",
    # Embeddings & RAG
    "EmbeddingClient", "CloudRuEmbeddingClient",
    "SparseEmbeddingClient",
    "VectorStore", "QdrantVectorStore",
    "Retriever",
    # Bulk indexation
    "ingest_and_index",
    # Query rewriters
    "RagQueryRewriter", "CrawlerQueryRewriter",
    # Internet
    "CrawlerClient",
    # LLM
    "OpenAIClient",
    # Orchestrator
    "ChatOrchestrator",
    # Factory
    "create_cloudru_openai_client_from_env",
    "create_cloudru_asr_client_from_env",
    "create_cloudru_embedding_client_from_env",
    "create_sparse_embedding_client_from_env",
    "create_default_orchestrator_from_env",
    # UC Utilities
    "generate_indexing_report",
    # UC Helpers (config, prompts, LLM wrappers)
    "load_env_and_validate",
    "create_rag_clients_from_env",
    "build_rag_prompt",
    "build_web_prompt",
    "stream_llm_to_stdout",
    "query_llm_simple",
    # Semantic Scholar
    "S2Client", "S2SearchFilter", "S2FieldInference",
    "S2_VALID_FIELDS_OF_STUDY", "S2_PAPER_FIELDS",
    "S2_AUTHOR_SEARCH_FIELDS", "S2_AUTHOR_DETAIL_FIELDS",
    "S2_CITATION_FIELDS", "S2_RATE_LIMIT_DELAY",
]